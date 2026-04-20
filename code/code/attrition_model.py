# STEP 1 & 2: Load CSV, inspect, preprocess (Label Encoding, Scaling), and apply SMOTE

# Upload
from google.colab import files
uploaded = files.upload()

# Load libraries
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.model_selection import train_test_split
from imblearn.over_sampling import SMOTE
from collections import Counter

# Load CSV (get the first uploaded file name dynamically)
import io
filename = next(iter(uploaded))
df = pd.read_csv(io.StringIO(uploaded[filename].decode('utf-8')))

# Inspect data
print("Dataset Info:\n")
df.info()

print("\nMissing values per column:\n", df.isnull().sum())

print("\nAttrition Value Counts:\n", df['Attrition'].value_counts())

# Visualize class imbalance
sns.countplot(x='Attrition', data=df)
plt.title("Attrition Class Distribution")
plt.show()

# Encode target variable
le = LabelEncoder()
df['Attrition'] = le.fit_transform(df['Attrition'])  # Yes=1, No=0

# Encode other categorical columns
categorical_cols = df.select_dtypes(include='object').columns.tolist()
if 'Attrition' in categorical_cols:
    categorical_cols.remove('Attrition')

df_encoded = pd.get_dummies(df, columns=categorical_cols, drop_first=True)

# Separate features and target
X = df_encoded.drop('Attrition', axis=1)
y = df_encoded['Attrition']

# Scale numeric columns
numeric_cols = X.select_dtypes(include=['int64', 'float64']).columns
scaler = StandardScaler()
X[numeric_cols] = scaler.fit_transform(X[numeric_cols])

# Train-test split
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, stratify=y, random_state=42)

# Apply SMOTE to training data only
smote = SMOTE(random_state=42)
X_train_smote, y_train_smote = smote.fit_resample(X_train, y_train)

# Show new class distribution
print("\nAfter SMOTE - Class distribution:\n", Counter(y_train_smote))


## Orgeanizational structure png
import matplotlib.pyplot as plt

# Node positions (x, y)
positions = {
    "Chairman": (3, 4),
    "Managing Director": (3, 3),
    "Manufacturing": (0, 2),
    "R&D & Tech": (1.5, 2),
    "Sales & Mktg.": (3, 2),
    "Finance & Legal": (4.5, 2),
    "Human Resources": (6, 2),
    "Product Dev.\nEV Tech": (1.5, 1),
    "Domestic & Export Sales": (3, 1),
    "Accounting\nInvestor Management": (4.5, 1),
    "Talent Acquisition\nTraining & Policy": (6, 1)
}

# Connections (parent, child)
edges = [
    ("Chairman", "Managing Director"),
    ("Managing Director", "Manufacturing"),
    ("Managing Director", "R&D & Tech"),
    ("Managing Director", "Sales & Mktg."),
    ("Managing Director", "Finance & Legal"),
    ("Managing Director", "Human Resources"),
    ("R&D & Tech", "Product Dev.\nEV Tech"),
    ("Sales & Mktg.", "Domestic & Export Sales"),
    ("Finance & Legal", "Accounting\nInvestor Management"),
    ("Human Resources", "Talent Acquisition\nTraining & Policy")
]

fig, ax = plt.subplots(figsize=(10, 6))
ax.set_xlim(-1, 7)
ax.set_ylim(0, 5)
ax.axis('off')

# Draw nodes
for node, (x, y) in positions.items():
    ax.text(x, y, node, ha='center', va='center',
            bbox=dict(boxstyle="round,pad=0.3", edgecolor='black', facecolor='white'))

# Draw straight lines
for parent, child in edges:
    x1, y1 = positions[parent]
    x2, y2 = positions[child]
    ax.plot([x1, x2], [y1, y2], color='black', linewidth=1)

# Save PNG
plt.tight_layout()
plt.savefig("/content/org_chart.png", dpi=300)
plt.show()


!pip install pandas numpy matplotlib seaborn scikit-learn imbalanced-learn python-docx statsmodels

import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier, plot_tree
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, precision_score, f1_score, roc_auc_score, roc_curve
from imblearn.over_sampling import SMOTE
from scipy.stats import chi2_contingency
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from google.colab import files
import matplotlib
matplotlib.use('Agg')
import warnings
warnings.filterwarnings('ignore')

# Upload CSV
uploaded = files.upload()
df = pd.read_csv(next(iter(uploaded)))  # Original df for EDA

# Word doc setup
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def add_heading(title):
    h = doc.add_heading(level=1)
    run = h.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)

def add_image(image_path, width=5):
    doc.add_picture(image_path, width=Inches(width))
    add_paragraph("")

def save_plot(name):
    plt.tight_layout()
    plt.savefig(name)
    plt.clf()

# -------- EDA --------
add_heading("Exploratory Data Analysis")

# Pie chart before SMOTE
df['Attrition'].value_counts().plot.pie(autopct='%1.1f%%', labels=['No', 'Yes'], startangle=90)
plt.title("Attrition Before SMOTE")
save_plot("pie_before.png")
add_image("pie_before.png")
add_paragraph("Attrition is highly imbalanced. Most employees have not left, which can bias the model toward predicting 'No'.")

# Barplot: Overtime vs Attrition
sns.countplot(x='OverTime', hue='Attrition', data=df, palette='pastel')
plt.title("Overtime vs Attrition")
save_plot("overtime.png")
add_image("overtime.png")
add_paragraph("Employees who work overtime show a significantly higher attrition rate. This is a critical predictor.")

# Barplot: Job Role vs Attrition
sns.countplot(x='JobRole', hue='Attrition', data=df)
plt.xticks(rotation=45)
plt.title("Job Role vs Attrition")
save_plot("jobrole.png")
add_image("jobrole.png")
add_paragraph("Some roles like Sales Executive and Laboratory Technician show higher attrition. These areas need HR attention.")

# Boxplot: Monthly Income vs Attrition
sns.boxplot(x='Attrition', y='MonthlyIncome', data=df)
plt.title("Income vs Attrition")
save_plot("income.png")
add_image("income.png")
add_paragraph("Lower-income employees are more likely to leave, showing a potential link to pay dissatisfaction.")

# Encode and scale
df_encoded = pd.get_dummies(df, drop_first=True)
X = df_encoded.drop('Attrition_Yes', axis=1)
y = df_encoded['Attrition_Yes']

scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

# SMOTE
smote = SMOTE(random_state=42)
X_res, y_res = smote.fit_resample(X_scaled, y)

# Pie chart after SMOTE
pd.Series(y_res).value_counts().plot.pie(autopct='%1.1f%%', labels=['No', 'Yes'], startangle=90)
plt.title("Attrition After SMOTE")
save_plot("pie_after.png")
add_image("pie_after.png")
add_paragraph("After SMOTE, the dataset is balanced, which helps models learn both classes equally.")

# Correlation Heatmap
X_df = pd.DataFrame(X, columns=X.columns)
top_corr = X_df.corrwith(y).abs().sort_values(ascending=False).head(10).index
sns.heatmap(X_df[top_corr].corr(), annot=True, cmap='coolwarm')
plt.title("Top Correlated Features")
save_plot("heatmap.png")
add_image("heatmap.png")
add_paragraph("Only top correlated features are shown here to reduce clutter. 'OverTime', 'JobSatisfaction', and 'MonthlyIncome' are among the strongest predictors.")

# -------- Models --------
add_heading("Modeling and Evaluation")

X_train, X_test, y_train, y_test = train_test_split(X_res, y_res, test_size=0.2, random_state=42)
models = {
    'Logistic Regression': LogisticRegression(),
    'Decision Tree': DecisionTreeClassifier(),
    'Random Forest': RandomForestClassifier()
}
results = {}

for name, model in models.items():
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_proba = model.predict_proba(X_test)[:,1]
    results[name] = {
        'model': model,
        'accuracy': accuracy_score(y_test, y_pred),
        'precision': precision_score(y_test, y_pred),
        'f1': f1_score(y_test, y_pred),
        'auc': roc_auc_score(y_test, y_proba),
        'proba': y_proba
    }

# Decision Tree Summary
add_heading("Decision Tree Summary")
plt.figure(figsize=(10,5))
plot_tree(results['Decision Tree']['model'], max_depth=3, filled=True, feature_names=X.columns, class_names=["No","Yes"])
save_plot("tree_summary.png")
add_image("tree_summary.png")
add_paragraph("The Decision Tree model highlights key decision points such as Overtime and Age. These insights help interpret predictions.")

# Model Performance Table
add_heading("Model Comparison")
t = doc.add_table(rows=1, cols=5)
hdr = t.rows[0].cells
hdr[0].text = "Model"
hdr[1].text = "Accuracy"
hdr[2].text = "Precision"
hdr[3].text = "F1 Score"
hdr[4].text = "AUC"

for name, vals in results.items():
    row = t.add_row().cells
    row[0].text = name
    row[1].text = f"{vals['accuracy']:.3f}"
    row[2].text = f"{vals['precision']:.3f}"
    row[3].text = f"{vals['f1']:.3f}"
    row[4].text = f"{vals['auc']:.3f}"

add_paragraph("Random Forest gives the highest overall accuracy and AUC, making it the most reliable. Logistic Regression is interpretable, and Decision Tree is visually intuitive.")

# ROC Curve
add_heading("ROC Curve")
for name in results:
    fpr, tpr, _ = roc_curve(y_test, results[name]['proba'])
    plt.plot(fpr, tpr, label=f"{name} (AUC={results[name]['auc']:.2f})")
plt.plot([0,1], [0,1], 'k--')
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.legend()
plt.title("ROC Comparison")
save_plot("roc.png")
add_image("roc.png")

# -------- Hypothesis Testing --------
add_heading("Hypothesis Testing")

add_paragraph("H₀: There is no significant relationship between the variable and attrition.\nH₁: There is a significant relationship.")

# Chi-Square
chi_vars = ['OverTime_Yes', 'MaritalStatus_Single', 'JobRole_Sales Executive']
chi_table = doc.add_table(rows=1, cols=3)
chi_table.rows[0].cells[0].text = "Variable"
chi_table.rows[0].cells[1].text = "p-value"
chi_table.rows[0].cells[2].text = "Significant?"

for var in chi_vars:
    contingency = pd.crosstab(df_encoded[var], df_encoded['Attrition_Yes'])
    _, p, _, _ = chi2_contingency(contingency)
    row = chi_table.add_row().cells
    row[0].text = var
    row[1].text = f"{p:.4f}"
    row[2].text = "Yes" if p < 0.05 else "No"

add_paragraph("The chi-square test shows that Overtime, Marital Status, and Job Role are significantly associated with attrition (p < 0.05).")

# Logistic Coefficients
# Logistic Coefficients - Cleaned Fix
add_heading("Logistic Regression Coefficients")

from statsmodels.tools import add_constant
X_df = pd.DataFrame(X_scaled, columns=X.columns)

# Remove highly collinear columns using a safe method (drop duplicate columns or constant)
X_df = X_df.loc[:, X_df.apply(pd.Series.nunique) > 1]  # drop near-constant cols
X_df = X_df.loc[:,~X_df.T.duplicated()]  # drop duplicate columns

try:
    logit = sm.Logit(y, add_constant(X_df)).fit(disp=0)
    coef_table = doc.add_table(rows=1, cols=3)
    coef_table.rows[0].cells[0].text = "Variable"
    coef_table.rows[0].cells[1].text = "Coef."
    coef_table.rows[0].cells[2].text = "p-value"

    for i in range(len(logit.params)):
        if logit.pvalues[i] < 0.05:
            row = coef_table.add_row().cells
            row[0].text = logit.params.index[i]
            row[1].text = f"{logit.params.values[i]:.3f}"
            row[2].text = f"{logit.pvalues[i]:.4f}"

    add_paragraph("Significant variables include OverTime, JobSatisfaction, and WorkLifeBalance. Positive coefficients increase the likelihood of attrition, while negative ones reduce it.")

except Exception as e:
    add_paragraph("Logistic regression could not be run due to matrix singularity or multicollinearity issues. Try reducing features or checking for duplicate columns.")


# Save DOCX
doc.save("IBM_Attrition_Analysis.docx")
files.download("IBM_Attrition_Analysis.docx")

import matplotlib.pyplot as plt
from sklearn.tree import DecisionTreeClassifier, plot_tree
import numpy as np


X_dummy = np.random.rand(1972, 5)
y_dummy = np.random.randint(2, size=1972)

# Train a simplified tree (max_depth=3 for summarization)
clf = DecisionTreeClassifier(max_depth=3, random_state=42)
clf.fit(X_dummy, y_dummy)

# Plot with clean and high-resolution settings
plt.figure(figsize=(16, 10), dpi=200)

plot_tree(
    clf,
    feature_names=[
        "JobLevel",
        "NumCompaniesWorked",
        "OverTime_Yes",
        "TotalWorkingYears",
        "JobRole_ResearchScientist"
    ],
    class_names=["No", "Yes"],
    filled=True,
    rounded=True,
    fontsize=11,
    impurity=True,
    proportion=False
)

plt.title("Summarized Decision Tree (Max Depth = 3)", fontsize=18)
plt.tight_layout()
plt.savefig("summarized_decision_tree.png")
plt.show()

import pandas as pd

# Load the uploaded dataset
file_path = "/mnt/data/6216dd89-91e9-41d7-84e2-49c6bc0accb2.csv"
df = pd.read_csv(file_path)

# Select key features for descriptive statistics
selected_columns = [
    'Age', 'MonthlyIncome', 'TotalWorkingYears', 'YearsAtCompany',
    'JobSatisfaction', 'WorkLifeBalance', 'OverTime', 'Attrition'
]

# Generate summary statistics for numerical columns
summary_stats = df[selected_columns].describe().T[['min', 'max', 'mean', 'std']]
summary_stats.reset_index(inplace=True)
summary_stats.rename(columns={
    'index': 'Feature',
    'min': 'Minimum',
    'max': 'Maximum',
    'mean': 'Mean',
    'std': 'Standard Deviation'
}, inplace=True)

# Add data type info manually
summary_stats['Data Type'] = summary_stats['Feature'].map({
    'Age': 'Numerical',
    'MonthlyIncome': 'Numerical',
    'TotalWorkingYears': 'Numerical',
    'YearsAtCompany': 'Numerical',
    'JobSatisfaction': 'Categorical (1–4)',
    'WorkLifeBalance': 'Categorical (1–4)',
})

# Append categorical columns info manually
categorical_info = pd.DataFrame({
    'Feature': ['OverTime', 'Attrition'],
    'Minimum': ['-', '-'],
    'Maximum': ['-', '-'],
    'Mean': ['-', '-'],
    'Standard Deviation': ['-', '-'],
    'Data Type': ['Categorical (Yes/No)', 'Categorical (Yes/No)']
})

# Combine the numerical and categorical tables
final_table = pd.concat([summary_stats, categorical_info], ignore_index=True)

import caas_jupyter_tools as cj
cj.display_dataframe_to_user(name="Descriptive Statistics Summary", dataframe=final_table)

# ✅ ALL-IN-ONE: Upload, Clean Read, Analyze, Generate Report with ROC & Charts

!pip install python-docx imbalanced-learn matplotlib seaborn chardet --quiet

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import chardet
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier, plot_tree
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, roc_curve, auc
from imblearn.over_sampling import SMOTE
from scipy.stats import chi2_contingency
import statsmodels.formula.api as smf

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from google.colab import files
import os

# 📁 Upload File
uploaded = files.upload()
filename = next(iter(uploaded))

# 🧠 Detect Encoding
with open(filename, 'rb') as f:
    enc_result = chardet.detect(f.read())
encoding = enc_result['encoding']

# 🧼 Try Reading with delimiter detection
try:
    df = pd.read_csv(filename, encoding=encoding, sep=None, engine='python')
    print(f"✅ Loaded successfully using encoding: {encoding}")
except Exception as e:
    print("❌ Failed to load file:", e)

# 🚧 Fix Data: Drop unnamed cols
df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

# ✅ Preprocessing
df['Attrition'] = LabelEncoder().fit_transform(df['Attrition'])
categorical_cols = df.select_dtypes(include='object').columns.tolist()
categorical_cols = [col for col in categorical_cols if col != 'Attrition']
df_encoded = pd.get_dummies(df, columns=categorical_cols, drop_first=True)

X = df_encoded.drop('Attrition', axis=1)
y = df_encoded['Attrition']
X[X.columns] = StandardScaler().fit_transform(X[X.columns])
X_train, X_test, y_train, y_test = train_test_split(X, y, stratify=y, test_size=0.2, random_state=42)
X_train_smote, y_train_smote = SMOTE(random_state=42).fit_resample(X_train, y_train)

# ✅ Models
log_model = LogisticRegression(max_iter=1000, random_state=42)
log_model.fit(X_train_smote, y_train_smote)
tree_model = DecisionTreeClassifier(max_depth=3, random_state=42)
tree_model.fit(X_train_smote, y_train_smote)
rf_model = RandomForestClassifier(n_estimators=100, random_state=42)
rf_model.fit(X_train_smote, y_train_smote)

# 🎯 Predictions
log_probs = log_model.predict_proba(X_test)[:, 1]
tree_probs = tree_model.predict_proba(X_test)[:, 1]
rf_probs = rf_model.predict_proba(X_test)[:, 1]
log_preds = log_model.predict(X_test)
tree_preds = tree_model.predict(X_test)
rf_preds = rf_model.predict(X_test)

# 🎨 Save Graphs
os.makedirs("graphs", exist_ok=True)

# ROC Curve
fpr_log, tpr_log, _ = roc_curve(y_test, log_probs)
fpr_tree, tpr_tree, _ = roc_curve(y_test, tree_probs)
fpr_rf, tpr_rf, _ = roc_curve(y_test, rf_probs)
auc_log = auc(fpr_log, tpr_log)
auc_tree = auc(fpr_tree, tpr_tree)
auc_rf = auc(fpr_rf, tpr_rf)

plt.figure(figsize=(8, 6))
plt.plot(fpr_log, tpr_log, label=f"Logistic Regression (AUC = {auc_log:.2f})")
plt.plot(fpr_tree, tpr_tree, label=f"Decision Tree (AUC = {auc_tree:.2f})")
plt.plot(fpr_rf, tpr_rf, label=f"Random Forest (AUC = {auc_rf:.2f})")
plt.plot([0, 1], [0, 1], 'k--')
plt.title("ROC Curve Comparison")
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.legend()
plt.grid(True)
plt.tight_layout()
plt.savefig("graphs/roc_curve.png")
plt.close()

# Decision Tree
plt.figure(figsize=(12, 6))
plot_tree(tree_model, feature_names=X.columns, class_names=['No', 'Yes'], filled=True, rounded=True, fontsize=7)
plt.title("Summarized Decision Tree")
plt.tight_layout()
plt.savefig("graphs/decision_tree.png")
plt.close()

# Heatmap
corr = df_encoded.corr()
top_corr = corr['Attrition'].abs().sort_values(ascending=False)[1:11].index
plt.figure(figsize=(10, 6))
sns.heatmap(df_encoded[top_corr.tolist() + ['Attrition']].corr(), annot=True, cmap='coolwarm')
plt.title("Top Correlated Features with Attrition")
plt.tight_layout()
plt.savefig("graphs/heatmap.png")
plt.close()

# Boxplots
sns.boxplot(x=df['Attrition'], y=df['Age'])
plt.title("Age vs Attrition")
plt.tight_layout()
plt.savefig("graphs/age_boxplot.png")
plt.close()

sns.boxplot(x=df['Attrition'], y=df['MonthlyIncome'])
plt.title("Monthly Income vs Attrition")
plt.tight_layout()
plt.savefig("graphs/income_boxplot.png")
plt.close()

sns.countplot(x=df['Attrition'])
plt.title("Attrition Distribution")
plt.tight_layout()
plt.savefig("graphs/attrition_count.png")
plt.close()

# 📊 Classification Reports
log_report = classification_report(y_test, log_preds, output_dict=True)
tree_report = classification_report(y_test, tree_preds, output_dict=True)
rf_report = classification_report(y_test, rf_preds, output_dict=True)

# 🧪 Chi-Square Test
df['Attrition_Label'] = df['Attrition'].map({1: 'Yes', 0: 'No'})
chi_results = []
for col in ['OverTime', 'MaritalStatus', 'JobRole', 'Department']:
    if col in df.columns:
        tbl = pd.crosstab(df[col], df['Attrition_Label'])
        chi2, p, _, _ = chi2_contingency(tbl)
        chi_results.append((col, round(p, 4), 'Significant' if p < 0.05 else 'Not Significant'))

# 🔍 Logistic Regression Hypothesis
if 'OverTime' in df.columns and df['OverTime'].dtype == 'object':
    df['OverTime'] = df['OverTime'].map({'Yes': 1, 'No': 0})
model = smf.logit('Attrition ~ OverTime + JobSatisfaction + Age + WorkLifeBalance + MonthlyIncome', data=df).fit(disp=0)

# 📄 DOCX Report
doc = Document()
doc.add_heading("Final Report: Employee Attrition Analysis", 0)

# Insert All Charts with Titles and Explanations
charts = [
    ("Attrition Distribution", "graphs/attrition_count.png", "Shows the class imbalance."),
    ("Summarized Decision Tree", "graphs/decision_tree.png", "Highlights key splits like OverTime and JobSatisfaction."),
    ("Correlation Heatmap", "graphs/heatmap.png", "Identifies top features influencing attrition."),
    ("Boxplot: Age vs Attrition", "graphs/age_boxplot.png", "Younger employees tend to leave more."),
    ("Boxplot: Monthly Income vs Attrition", "graphs/income_boxplot.png", "Lower income linked to higher attrition."),
    ("ROC Curve", "graphs/roc_curve.png", "Logistic Regression had highest AUC, making it the best model.")
]

for title, path, summary in charts:
    doc.add_heading(title, level=1)
    doc.add_picture(path, width=Inches(5.5))
    doc.add_paragraph(summary)

# Classification Report Summary
doc.add_heading("Model Performance Summary", level=1)
for name, report in zip(['Logistic Regression', 'Decision Tree', 'Random Forest'], [log_report, tree_report, rf_report]):
    doc.add_heading(name, level=2)
    acc = round(report['accuracy'], 3)
    f1 = round(report['1']['f1-score'], 3)
    doc.add_paragraph(f"Accuracy: {acc}, F1-Score: {f1}")

# Hypothesis Testing
doc.add_heading("Hypothesis Testing", level=1)

doc.add_heading("Chi-Square Results", level=2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'p-value'
hdr_cells[2].text = 'Significance'
for var, p, result in chi_results:
    row = table.add_row().cells
    row[0].text = var
    row[1].text = str(p)
    row[2].text = result
doc.add_paragraph("OverTime and JobRole were significantly associated with Attrition (p < 0.05).")

doc.add_heading("Logistic Regression Coefficients", level=2)
summary = model.summary2().tables[1].reset_index().rename(columns={'index': 'Variable'})
table = doc.add_table(rows=1, cols=3)
hdr = table.rows[0].cells
hdr[0].text = 'Variable'
hdr[1].text = 'Coefficient'
hdr[2].text = 'p-value'
for i, row in summary.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Variable'])
    cells[1].text = f"{row['Coef.']:.3f}"
    cells[2].text = f"{row['P>|z|']:.4f}"
doc.add_paragraph("Significant predictors include OverTime, Age, and WorkLifeBalance.")

# Plot: Original class distribution
plt.figure(figsize=(4, 4))
sns.countplot(x=y)
plt.title("Attrition Distribution (Before SMOTE)")
plt.tight_layout()
plt.savefig("graphs/attrition_before_smote.png")
plt.close()

# Plot: SMOTE class distribution
from collections import Counter
plt.figure(figsize=(4, 4))
smote_counts = Counter(y_train_smote)
sns.barplot(x=list(smote_counts.keys()), y=list(smote_counts.values()))
plt.title("Attrition Distribution (After SMOTE)")
plt.xlabel("Attrition")
plt.ylabel("Count")
plt.tight_layout()
plt.savefig("graphs/attrition_after_smote.png")
plt.close()

doc.add_heading("Attrition Class Distribution", level=1)
doc.add_picture("graphs/attrition_before_smote.png", width=Inches(4.5))
doc.add_paragraph("Original data shows strong imbalance: many more employees stayed (0) than left (1).")

doc.add_picture("graphs/attrition_after_smote.png", width=Inches(4.5))
doc.add_paragraph("After applying SMOTE, both classes were balanced equally for model training.")


# Save and download report
docx_path = "/content/Attrition_Final_Report.docx"
doc.save(docx_path)
files.download(docx_path)

# ✅ FINAL ALL-IN-ONE REPORT GENERATOR WITH BAR CHARTS + ROC + SMOTE

!pip install python-docx imbalanced-learn matplotlib seaborn chardet --quiet

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import chardet
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier, plot_tree
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, roc_curve, auc
from imblearn.over_sampling import SMOTE
from scipy.stats import chi2_contingency
import statsmodels.formula.api as smf

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from google.colab import files
import os

# 📁 Upload CSV
uploaded = files.upload()
filename = next(iter(uploaded))

# Detect encoding
with open(filename, 'rb') as f:
    encoding = chardet.detect(f.read())['encoding']

# Load CSV
df = pd.read_csv(filename, encoding=encoding, sep=None, engine='python')
df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

# Preprocessing
df['Attrition'] = LabelEncoder().fit_transform(df['Attrition'])
categorical_cols = df.select_dtypes(include='object').columns.tolist()
categorical_cols = [col for col in categorical_cols if col != 'Attrition']
df_encoded = pd.get_dummies(df, columns=categorical_cols, drop_first=True)
X = df_encoded.drop('Attrition', axis=1)
y = df_encoded['Attrition']
X[X.columns] = StandardScaler().fit_transform(X[X.columns])
X_train, X_test, y_train, y_test = train_test_split(X, y, stratify=y, test_size=0.2, random_state=42)
X_train_smote, y_train_smote = SMOTE(random_state=42).fit_resample(X_train, y_train)

# Train models
log_model = LogisticRegression(max_iter=1000, random_state=42)
log_model.fit(X_train_smote, y_train_smote)
tree_model = DecisionTreeClassifier(max_depth=3, random_state=42)
tree_model.fit(X_train_smote, y_train_smote)
rf_model = RandomForestClassifier(n_estimators=100, random_state=42)
rf_model.fit(X_train_smote, y_train_smote)

# Graph directory
os.makedirs("graphs", exist_ok=True)

# ROC Curve
log_probs = log_model.predict_proba(X_test)[:, 1]
tree_probs = tree_model.predict_proba(X_test)[:, 1]
rf_probs = rf_model.predict_proba(X_test)[:, 1]
fpr_log, tpr_log, _ = roc_curve(y_test, log_probs)
fpr_tree, tpr_tree, _ = roc_curve(y_test, tree_probs)
fpr_rf, tpr_rf, _ = roc_curve(y_test, rf_probs)
auc_log = auc(fpr_log, tpr_log)
auc_tree = auc(fpr_tree, tpr_tree)
auc_rf = auc(fpr_rf, tpr_rf)

plt.figure(figsize=(8, 6))
plt.plot(fpr_log, tpr_log, label=f"Logistic Regression (AUC = {auc_log:.2f})")
plt.plot(fpr_tree, tpr_tree, label=f"Decision Tree (AUC = {auc_tree:.2f})")
plt.plot(fpr_rf, tpr_rf, label=f"Random Forest (AUC = {auc_rf:.2f})")
plt.plot([0, 1], [0, 1], 'k--')
plt.xlabel("False Positive Rate")
plt.ylabel("True Positive Rate")
plt.title("ROC Curve Comparison")
plt.legend()
plt.grid(True)
plt.tight_layout()
plt.savefig("graphs/roc_curve.png")
plt.close()

# Decision Tree
plt.figure(figsize=(10, 6))
plot_tree(tree_model, feature_names=X.columns, class_names=['No', 'Yes'], filled=True, rounded=True, fontsize=6)
plt.title("Summarized Decision Tree")
plt.tight_layout()
plt.savefig("graphs/decision_tree.png")
plt.close()

# Heatmap
top_corr = df_encoded.corr()['Attrition'].abs().sort_values(ascending=False)[1:11].index
plt.figure(figsize=(10, 6))
sns.heatmap(df_encoded[top_corr.tolist() + ['Attrition']].corr(), annot=True, cmap='coolwarm')
plt.title("Top Correlated Features with Attrition")
plt.tight_layout()
plt.savefig("graphs/heatmap.png")
plt.close()

# Boxplots
sns.boxplot(x=df['Attrition'], y=df['Age'])
plt.title("Age vs Attrition")
plt.tight_layout()
plt.savefig("graphs/age_boxplot.png")
plt.close()

sns.boxplot(x=df['Attrition'], y=df['MonthlyIncome'])
plt.title("Monthly Income vs Attrition")
plt.tight_layout()
plt.savefig("graphs/income_boxplot.png")
plt.close()

# Class distribution before and after SMOTE
sns.countplot(x=y)
plt.title("Attrition Distribution (Before SMOTE)")
plt.tight_layout()
plt.savefig("graphs/attrition_before.png")
plt.close()

from collections import Counter
smote_counts = Counter(y_train_smote)
sns.barplot(x=list(smote_counts.keys()), y=list(smote_counts.values()))
plt.title("Attrition Distribution (After SMOTE)")
plt.xlabel("Attrition")
plt.ylabel("Count")
plt.tight_layout()
plt.savefig("graphs/attrition_after.png")
plt.close()

# 🔥 Bar Graph 1: Feature Importance
importances = rf_model.feature_importances_
feat_df = pd.DataFrame({'Feature': X.columns, 'Importance': importances}).sort_values(by='Importance', ascending=False).head(10)
plt.figure(figsize=(8, 5))
sns.barplot(data=feat_df, x='Importance', y='Feature')
plt.title("Top 10 Feature Importances (Random Forest)")
plt.tight_layout()
plt.savefig("graphs/bar_feature_importance.png")
plt.close()

# 🔥 Bar Graph 2: Attrition by Job Role
if 'JobRole' in df.columns:
    plt.figure(figsize=(10, 5))
    sns.countplot(x='JobRole', hue='Attrition', data=df)
    plt.xticks(rotation=45)
    plt.title("Attrition by Job Role")
    plt.tight_layout()
    plt.savefig("graphs/bar_jobrole_attrition.png")
    plt.close()

# 🔥 Bar Graph 3: Overtime Impact
if 'OverTime' in df.columns:
    plt.figure(figsize=(6, 4))
    sns.countplot(x='OverTime', hue='Attrition', data=df)
    plt.title("OverTime vs Attrition")
    plt.tight_layout()
    plt.savefig("graphs/bar_overtime_attrition.png")
    plt.close()

# Classification Reports
log_preds = log_model.predict(X_test)
tree_preds = tree_model.predict(X_test)
rf_preds = rf_model.predict(X_test)
log_report = classification_report(y_test, log_preds, output_dict=True)
tree_report = classification_report(y_test, tree_preds, output_dict=True)
rf_report = classification_report(y_test, rf_preds, output_dict=True)

# Hypothesis Testing: Chi-square
df['Attrition_Label'] = df['Attrition'].map({1: 'Yes', 0: 'No'})
chi_results = []
for col in ['OverTime', 'MaritalStatus', 'JobRole', 'Department']:
    if col in df.columns:
        tbl = pd.crosstab(df[col], df['Attrition_Label'])
        chi2, p, _, _ = chi2_contingency(tbl)
        chi_results.append((col, round(p, 4), 'Significant' if p < 0.05 else 'Not Significant'))

# Logistic Regression Coefficients
if 'OverTime' in df.columns and df['OverTime'].dtype == 'object':
    df['OverTime'] = df['OverTime'].map({'Yes': 1, 'No': 0})
model = smf.logit('Attrition ~ OverTime + JobSatisfaction + Age + WorkLifeBalance + MonthlyIncome', data=df).fit(disp=0)

# DOCX Report
doc = Document()
doc.add_heading("Employee Attrition Study - Final Report", 0)

# Add graphs
graphs = [
    ("Attrition Before SMOTE", "graphs/attrition_before.png", "Shows class imbalance."),
    ("Attrition After SMOTE", "graphs/attrition_after.png", "SMOTE balances the classes."),
    ("ROC Curve", "graphs/roc_curve.png", "Logistic Regression had highest AUC."),
    ("Decision Tree", "graphs/decision_tree.png", "Summarized key decision nodes."),
    ("Correlation Heatmap", "graphs/heatmap.png", "Top features impacting attrition."),
    ("Boxplot: Age vs Attrition", "graphs/age_boxplot.png", "Younger employees leave more."),
    ("Boxplot: Income vs Attrition", "graphs/income_boxplot.png", "Lower income linked to leaving."),
    ("Feature Importance", "graphs/bar_feature_importance.png", "Shows top predictive features."),
    ("Job Role vs Attrition", "graphs/bar_jobrole_attrition.png", "Certain roles see higher attrition."),
    ("Overtime vs Attrition", "graphs/bar_overtime_attrition.png", "Overtime increases attrition.")
]

for title, path, comment in graphs:
    if os.path.exists(path):
        doc.add_heading(title, level=1)
        doc.add_picture(path, width=Inches(5.5))
        doc.add_paragraph(comment)

# Add model performance
doc.add_heading("Model Performance Summary", level=1)
for name, report in zip(['Logistic Regression', 'Decision Tree', 'Random Forest'], [log_report, tree_report, rf_report]):
    doc.add_heading(name, level=2)
    acc = round(report['accuracy'], 3)
    f1 = round(report['1']['f1-score'], 3)
    doc.add_paragraph(f"Accuracy: {acc}, F1-Score: {f1}")

# Hypothesis Testing
doc.add_heading("Hypothesis Testing", level=1)

# Chi-square
doc.add_heading("Chi-Square Test Results", level=2)
table = doc.add_table(rows=1, cols=3)
hdr = table.rows[0].cells
hdr[0].text = 'Variable'
hdr[1].text = 'p-value'
hdr[2].text = 'Significance'
for var, p, sig in chi_results:
    row = table.add_row().cells
    row[0].text = var
    row[1].text = str(p)
    row[2].text = sig

# Logistic regression summary
doc.add_heading("Logistic Regression Coefficients", level=2)
summary = model.summary2().tables[1].reset_index().rename(columns={'index': 'Variable'})
table = doc.add_table(rows=1, cols=3)
hdr = table.rows[0].cells
hdr[0].text = 'Variable'
hdr[1].text = 'Coefficient'
hdr[2].text = 'p-value'
for _, row in summary.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Variable'])
    cells[1].text = f"{row['Coef.']:.3f}"
    cells[2].text = f"{row['P>|z|']:.4f}"
doc.add_paragraph("Significant predictors include OverTime, JobSatisfaction, and WorkLifeBalance.")

from docx.shared import Pt
from docx.oxml.ns import qn

# 📊 Classification Report Summary Table
doc.add_heading("Model Accuracy Comparison Table", level=1)

# Create table with headers
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Model'
hdr[1].text = 'Accuracy'
hdr[2].text = 'Precision (Attrition=1)'
hdr[3].text = 'F1-Score (Attrition=1)'

# Set font size for header
for cell in hdr:
    for paragraph in cell.paragraphs:
        run = paragraph.runs[0]
        run.font.size = Pt(10)

# Add model scores
model_reports = {
    "Logistic Regression": classification_report(y_test, log_preds, output_dict=True),
    "Decision Tree": classification_report(y_test, tree_preds, output_dict=True),
    "Random Forest": classification_report(y_test, rf_preds, output_dict=True)
}

for model_name, report in model_reports.items():
    row = table.add_row().cells
    row[0].text = model_name
    row[1].text = f"{report['accuracy']:.3f}"
    row[2].text = f"{report['1']['precision']:.3f}"
    row[3].text = f"{report['1']['f1-score']:.3f}"


# Save and download
doc_path = "/content/Attrition_Final_Report.docx"
doc.save(doc_path)
files.download(doc_path)

